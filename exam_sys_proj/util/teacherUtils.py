from ..dao.HomeworkOrExamPoolDAO import HomeworkOrExamPool, HomeworkOrExamPoolDAO
from ..dao.KnowledgePointsDAO import KnowledgePointsDAO, KnowledgePoints
from ..dao.HepAndKpMediaterDAO import HepAndKpMediater, HepAndKpMediaterDAO
from ..dao.TeacherCourseDAO import TeacherCourse, TeacherCourseDAO
from collections import defaultdict
import markdown2
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
import json
import random
import time
import copy


class TeacherUtils:

    def __init__(self):
        self.a = None

    @classmethod
    def insertOneQuestion(cls, db_util, question_dict: dict):
        """
        传入前端生成的题目字典，将题目存入数据库\n
        并没有保证插入的原子性，无法保证数据的完整性
        :param db_util:
        :param question_dict:
        :return:
        """
        homeworkOrExamPoolDao = HomeworkOrExamPoolDAO(db_util)
        hepAndKpMediaterDAO = HepAndKpMediaterDAO(db_util)
        homeworkOrExamPool = HomeworkOrExamPool(
            type=question_dict["type"],
            answer=question_dict["answer"],
            courseName=question_dict['subject'],
            difficultyLevel=question_dict['difficulty'],
            isActive=True
        )
        kp_list = []
        if question_dict['knowledge_point']:
            for i in question_dict['knowledge_point']:
                kp_list.append(i)
            del question_dict["answer"]
            del question_dict['knowledge_point']
            del question_dict['difficulty']
            del question_dict['subject']
            homeworkOrExamPool.question = question_dict
            hepID = homeworkOrExamPoolDao.insert(homeworkOrExamPool)
            if len(kp_list) > 0:
                for i in kp_list:
                    hepAndKpMediater = HepAndKpMediater(hepID=hepID, kpName=i)
                    hepAndKpMediaterDAO.insert(hepAndKpMediater)
        else:
            del question_dict["answer"]
            del question_dict['knowledge_point']
            del question_dict['difficulty']
            del question_dict['subject']
            homeworkOrExamPool.question = question_dict
            homeworkOrExamPoolDao.insert(homeworkOrExamPool)

    @classmethod
    def queryTeacherSubjectKP(cls, db_util, userID: int):
        """
        通过教师的id查询教师所教授的学科的知识点
        传入dbutil初始化，传入教师的userID进行查询，返回一个元素是“KnowledgePoints类的orm对象”的list
        :param db_util:
        :param userID:
        :return : 元素是knowledgePoints的orm的list
        """
        teacherCourseDAO = TeacherCourseDAO(db_util)
        knowledgePointsDAO = KnowledgePointsDAO(db_util)
        subject_list = teacherCourseDAO.querySubjectViaTeacherID(userID)
        result_list = knowledgePointsDAO.query(value=subject_list[0].subject, column_name="subject", is_all=True)
        return result_list

    @classmethod
    def batchInsertQuestions(cls, db_util, json_path: str):
        """
        这是批量插入题目的丐版；并没有保证插入的原子性，无法保证数据的完整性
        :param db_util:
        :param json_path: json的文件路径
        :return:
        """
        try:
            question_list = cls._readOurJson(json_path=json_path)
            homeworkOrExamPoolDao = HomeworkOrExamPoolDAO(db_util)
            hepAndKpMediaterDAO = HepAndKpMediaterDAO(db_util)
            homeworkOrExamPool_list = []
            kp_list_all = []
            for question_dict in question_list:
                homeworkOrExamPool = HomeworkOrExamPool(
                    type=question_dict["type"],
                    answer=question_dict["answer"],
                    courseName=question_dict['subject'],
                    difficultyLevel=question_dict['difficulty'],
                    isActive=True
                )
                kp_list = []
                for i in question_dict['knowledge_point']:
                    kp_list.append(i)
                del question_dict["answer"]
                del question_dict['knowledge_point']
                del question_dict['difficulty']
                del question_dict['subject']
                homeworkOrExamPool.question = question_dict
                homeworkOrExamPool_list.append(homeworkOrExamPool)
                kp_list_all.append(kp_list)
            pk_list = homeworkOrExamPoolDao.batchInsert(homeworkOrExamPool_list)
            hepAndKpMediater_list = []
            for i in range(len(kp_list_all)):
                if len(kp_list_all[i]) > 0:
                    for j in kp_list_all[i]:
                        hepAndKpMediater = HepAndKpMediater(hepID=pk_list[i], kpName=j)
                        hepAndKpMediater_list.append(hepAndKpMediater)
            hepAndKpMediaterDAO.batchInsert(hepAndKpMediater_list)
        except Exception as e:
            print(e)
            print("batchInsertQuestions error!")
            return "error"

    @classmethod
    def random_paper(cls, input_dict: dict, db_util):
        random.seed(time.time())
        # 处理输入数据，删除题目数量为0的大题, 和数量为0的知识点
        temp_list0 = []
        for _ in ['选择题', '判断题', '填空题', '主观题']:
            temp_sum0 = 0
            for __ in input_dict[_]["amount_per_knowledge_point"]:
                temp_sum0 += eval(input_dict[_]["amount_per_knowledge_point"][__])
            if temp_sum0 == 0:
                del input_dict[_]
            else:
                temp_list0.append(_)
                del_kp_list = []
                for __ in input_dict[_]["amount_per_knowledge_point"]:
                    if input_dict[_]["amount_per_knowledge_point"][__] == "0":
                        del_kp_list.append(__)
                for __ in del_kp_list:
                    del input_dict[_]["amount_per_knowledge_point"][__]
        if len(temp_list0) == 0:
            return {"status_code": 404, "message": "试卷里至少要有一道题",
                    "content": None}

        def get_diff(min_d, max_d):
            min_d = eval(min_d)
            max_d = eval(max_d)
            if 1 <= min_d <= 3 and 1 <= max_d <= 3:
                return 1,
            elif 1 <= min_d <= 3 and 4 <= max_d <= 7:
                return 1, 2
            elif 1 <= min_d <= 3 and 8 <= max_d <= 10:
                return 1, 2, 3
            elif 4 <= min_d <= 7 and 8 <= max_d <= 10:
                return 2, 3
            elif 4 <= min_d <= 7 and 4 <= max_d <= 7:
                return 2,
            else:
                return 3,

        def get_interval(diff, min_d, max_d):
            set1 = set()
            if diff == 1:
                set1 = set(range(1, 4))
            elif diff == 2:
                set1 = set(range(4, 8))
            elif diff == 3:
                set1 = set(range(8, 11))
            return list(set1 & set(range(eval(min_d), eval(max_d) + 1)))

        questions_dict = dict()
        score_dict = dict()
        score_ratio = [
            eval(input_dict['overall_difficulty_easy']) / (
                    eval(input_dict['overall_difficulty_easy']) + eval(input_dict['overall_difficulty_normal']) + eval(
                input_dict['overall_difficulty_hard'])),
            eval(input_dict['overall_difficulty_normal']) / (
                    eval(input_dict['overall_difficulty_easy']) + eval(input_dict['overall_difficulty_normal']) + eval(
                input_dict['overall_difficulty_hard'])),
            eval(input_dict['overall_difficulty_hard']) / (
                    eval(input_dict['overall_difficulty_easy']) + eval(input_dict['overall_difficulty_normal']) + eval(
                input_dict['overall_difficulty_hard']))
        ]

        for _ in temp_list0:
            questions_dict[_] = input_dict[_].copy()
            score_dict[_] = 0
        questions_type_list = temp_list0.copy()
        del temp_list0
        all_sum_score = 0
        for i in questions_dict:
            sum_score = 0
            sum_count = 0
            base_score = questions_dict[i]['score_per_question']
            if base_score == "0.0":
                base_score = "4"
            elif base_score == "0":
                base_score = "15"
            base_score = eval(base_score)
            diff_list = get_diff(questions_dict[i]['difficulty_min'], questions_dict[i]['difficulty_max'])
            for j in questions_dict[i]["amount_per_knowledge_point"]:
                sum_score += eval(questions_dict[i]["amount_per_knowledge_point"][j]) * base_score
                sum_count += eval(questions_dict[i]["amount_per_knowledge_point"][j])
                score_dict[i] = [diff_list, sum_score, sum_count]
            all_sum_score += sum_score
        score_list = sorted(score_dict.items(), key=lambda x: len(x[1][0]))
        sum_list = list()
        for i in score_ratio:
            sum_list.append(i * all_sum_score)
        # print(sum_list)
        over_list = [-0.15 * sum_list[0], -0.15 * sum_list[1], -0.15 * sum_list[2]]
        # print(over_list, "over_list")
        questions_diff_count_dict = dict()
        for i in score_list:
            questions_diff_count_dict[i[0]] = list()
            temp_sum = sum(sum_list[i[1][0][0] - 1:i[1][0][0] - 1 + len(i[1][0])])
            for j in i[1][0]:
                temp_rite = (sum_list[j - 1] / temp_sum)
                sum_list[j - 1] -= i[1][1] * temp_rite
                questions_diff_count_dict[i[0]] += [j, i[1][2] * temp_rite]
            if sum_list[0] < over_list[0]:
                print("error")
                return {"status_code": 404, "message": "所设定简单题数量过多，请调高大题的难度范围，或调高简单题所占比例",
                        "content": None}
            elif sum_list[1] < over_list[1]:
                print("error")
                return {"status_code": 404, "message": "所设定中等题数量过多，请调高大题的难度范围，或调高中等题所占比例",
                        "content": None}
            elif sum_list[2] < over_list[2]:
                print("error")
                return {"status_code": 404, "message": "所设定困难题数量过多，请调高大题的难度范围，或调高困难题所占比例",
                        "content": None}
        # del temp_sum, temp_rite
        all_query_kp_dict = dict()
        for i in questions_diff_count_dict:
            final_dict = defaultdict(int)
            # 舍入规则：小于1一律为1，其余四舍五入，若与总数不符则随机差补
            for j in range(1, len(questions_diff_count_dict[i]), 2):
                if questions_diff_count_dict[i][j] < 1:
                    questions_diff_count_dict[i][j] = 1
                else:
                    questions_diff_count_dict[i][j] = round(questions_diff_count_dict[i][j])
            if sum(questions_diff_count_dict[i][1::2]) > score_dict[i][2]:
                temp_n = 0
                temp_n2 = 20
                while temp_n == 0:
                    if questions_diff_count_dict[i][random.randrange(1, len(questions_diff_count_dict[i]), 2)] == 1:
                        pass
                    else:
                        questions_diff_count_dict[i][random.randrange(1, len(questions_diff_count_dict[i]), 2)] -= 1
                        temp_n = 1
                    if temp_n2 == 0:
                        temp_n = 1
                    temp_n2 -= 1
            elif sum(questions_diff_count_dict[i][1::2]) < score_dict[i][2]:
                questions_diff_count_dict[i][random.randrange(1, len(questions_diff_count_dict[i]), 2)] += 1
            for j in range(0, len(questions_diff_count_dict[i]), 2):
                questions_diff_count_dict[i][j] = get_interval(questions_diff_count_dict[i][j],
                                                               questions_dict[i]['difficulty_min'],
                                                               questions_dict[i]['difficulty_max'])
            diff_count_list = []
            for j in range(1, len(questions_diff_count_dict[i]), 2):
                for _ in range(questions_diff_count_dict[i][j]):
                    diff_count_list.append(random.choice(questions_diff_count_dict[i][j - 1]))
            for j in diff_count_list:
                final_dict[j] += 1
            questions_diff_count_dict[i] = dict(final_dict)

            temp_diff_list = list(questions_diff_count_dict[i].keys())
            if input_dict["shuffle"]:
                random.shuffle(temp_diff_list)
            else:
                temp_diff_list = sorted(temp_diff_list)
            temp_dict = dict()
            for j in temp_diff_list:
                temp_dict[j] = questions_diff_count_dict[i][j]
            # 得到难度数量字典
            questions_diff_count_dict[i] = temp_dict
            # 得到查询知识点字典
            query_kp_dict = cls._query_diff_to_kp(i, temp_diff_list, db_util)
            # 检查是否存在某个难度的题目数量为0

            if isinstance(query_kp_dict, dict):
                pass
            else:
                return {"status_code": 404,
                        "message": f"组卷失败，试题库中难度为{query_kp_dict}的{i}数量为零，若您是老师或管理员请在试题库中添加题目后再重试",
                        "content": None}
            all_query_kp_dict[i] = query_kp_dict
            del temp_dict, temp_diff_list

        questionID_dict = dict()
        kp_list_dict = dict()

        for question_type in questions_diff_count_dict:
            # 初始化input_dict
            for _ in input_dict[question_type]["amount_per_knowledge_point"]:
                input_dict[question_type]["amount_per_knowledge_point"][_] = int(
                    input_dict[question_type]["amount_per_knowledge_point"][_])
            # 使用analysis_diff_to_kp时允许3次失误
            temp_set = set()
            questionID_list = []
            kp_list = []
            # 数据检查
            temp_key1, temp_key2 = questions_diff_count_dict[question_type].keys(), input_dict[question_type][
                "amount_per_knowledge_point"].keys()
            if sum(questions_diff_count_dict[question_type].values()) - sum(
                    input_dict[question_type]["amount_per_knowledge_point"].values()) == 1:
                questions_diff_count_dict[question_type][min(questions_diff_count_dict[question_type].keys())] -= 1
                if questions_diff_count_dict[question_type][min(questions_diff_count_dict[question_type].keys())] == 0:
                    del questions_diff_count_dict[question_type][min(questions_diff_count_dict[question_type].keys())]
            elif sum(questions_diff_count_dict[question_type].values()) - sum(
                    input_dict[question_type]["amount_per_knowledge_point"].values()) == 2:
                questions_diff_count_dict[question_type][min(questions_diff_count_dict[question_type].keys())] -= 1
                if questions_diff_count_dict[question_type][min(questions_diff_count_dict[question_type].keys())] == 0:
                    del questions_diff_count_dict[question_type][min(questions_diff_count_dict[question_type].keys())]
                questions_diff_count_dict[question_type][max(questions_diff_count_dict[question_type].keys())] -= 1
                if questions_diff_count_dict[question_type][max(questions_diff_count_dict[question_type].keys())] == 0:
                    del questions_diff_count_dict[question_type][max(questions_diff_count_dict[question_type].keys())]

            # 判断是否有足够的查询数据能支撑难度的数量
            # 比如难度为1的题我们需要3条但是我们查寻出的难度为1的题目只有两道就不满足
            for d_l in questions_diff_count_dict[question_type]:
                if questions_diff_count_dict[question_type][d_l] > len(all_query_kp_dict[question_type][d_l]):
                    return {"status_code": 404,
                            "message": f"组卷失败，试题库中难度为{d_l}的{question_type}数量不足，若您是老师或管理员请在试题库中添加题目后再重试",
                            "content": None}

            for _ in range(3):
                try:

                    questionID_list, kp_list = cls._analysis_diff_to_kp(questions_diff_count_dict[question_type],
                                                                        input_dict[question_type][
                                                                            "amount_per_knowledge_point"],
                                                                        all_query_kp_dict[question_type])
                except Exception as e:
                    print("break1")
                    print(e)
                    return {"status_code": 404,
                            "message": "组卷失败，发生未知错误",
                            "content": None}

                # questionID_list, kp_list = cls._analysis_diff_to_kp(questions_diff_count_dict[question_type],
                #                                                input_dict[question_type]["amount_per_knowledge_point"],
                #                                                all_query_kp_dict[question_type])

                if isinstance(questionID_list, list):
                    break
                else:
                    temp_set = temp_set | questionID_list
                if _ == 2:
                    print(temp_set)
                    if len(temp_set) > 1:
                        return {"status_code": 404,
                                "message": f"组卷失败，试题库中知识点为{'、'.join(list(temp_set))}的{question_type}数量不足，若您是老师或管理员请在试题库中添加题目后再重试",
                                "content": None}
                    else:
                        return {"status_code": 404,
                                "message": f"组卷失败，试题库中知识点为{list(temp_set)[0]}的{question_type}数量不足，若您是老师或管理员请在试题库中添加题目后再重试",
                                "content": None}
            questionID_list.insert(0, eval(questions_dict[question_type]['score_per_question']))
            questionID_dict[question_type] = questionID_list
            kp_list_dict[question_type] = kp_list
        stored_paper_dict = dict()
        stored_paper_dict["type"] = input_dict["type"]
        stored_paper_dict["score"] = []
        stored_paper_dict["shuffle"] = input_dict["shuffle"]
        stored_paper_dict["main_content"] = input_dict["title"]
        stored_paper_dict["questions"] = questionID_dict
        try:
            homeworkOrExamPoolDAO = HomeworkOrExamPoolDAO(db_util)
            result = homeworkOrExamPoolDAO.query_whole_paper(stored_paper_dict)
            stored_paper_dict["score"] = result[0]["score"]
            stored_paper_dict["answer"] = result[1]
            stored_paper_dict["subject"] = input_dict["subject"]
            stored_paper_dict["difficulty"] = int(sum(result[2]) / len(result[2]))
            # kp_set = set()
            # for _ in kp_list_dict:
            #     for __ in kp_list_dict[_]:
            #         for ___ in __:
            #             kp_set.add(___)
            # kp_set = list(kp_set)
            stored_paper_dict["knowledge_point"] = None
        #     可能后续会有功能拓展
        except Exception as e:
            print(e)
            return {"status_code": 404,
                    "message": "组卷失败，发生未知错误",
                    "content": None}

        # 分析试卷
        analysis_text = cls._analysis_paper(stored_paper_dict, result[2], questions_type_list, None)
        # 输出
        # print(questions_dict)

        return {"status_code": 200,
                "message": "组卷成功",
                "content": [result[0], stored_paper_dict, analysis_text, result[1]]}

    @classmethod
    def querySubjectQuestionsViaUID(cls, db_util, userID, entity: HomeworkOrExamPool):
        """
        查询试题\n
        示例说明：\n
        有一个：HomeworkOrExamPool对象：homeworkOrExamPool\n
        userid = 24\n
        homeworkOrExamPool.type = "选择题"\n
        homeworkOrExamPool.difficultyLevel = 5\n
        将userid和homeworkOrExamPool传入方法，查询得到24号老师所教学科的所有难度为5的选择题
        :param db_util: 数据库连接池对象
        :param userID: 教师的userid
        :param entity: HomeworkOrExamPool对象
        :return: 元素是HomeworkOrExamPool对象的列表
        """
        teacherCourseDAO = TeacherCourseDAO(db_util)
        homeworkOrExamPoolDAO = HomeworkOrExamPoolDAO(db_util)
        subject_list = teacherCourseDAO.querySubjectViaTeacherID(userID)
        entity.courseName = subject_list[0].subject
        result_list = homeworkOrExamPoolDAO.columnsQuery(entity, is_all=True)
        return result_list

    @classmethod
    def export_paper_as_docx(cls, db_util, hepID: int, path: str):
        homeworkOrExamPoolDAO = HomeworkOrExamPoolDAO(db_util)
        entity: HomeworkOrExamPool = homeworkOrExamPoolDAO.query(hepID)
        whole_paper = homeworkOrExamPoolDAO.query_whole_paper(json.loads(entity.question))
        md: str = whole_paper[0]["main_content"]
        md.replace("<center>", "")
        md.replace("</center>", "")
        for i in whole_paper[0]["questions"]:
            if isinstance(i, str):
                md += i
                md += "\n"
            elif isinstance(i, dict):
                md += i["main_content"]
                md += "\n"
                for j in i["questions"]:
                    md += j
                    md += "\n"
                md += "\n"
        cls._markdown_to_word(md, path)

    @classmethod
    def export_questions(cls, db_util, subject: str, path: str):
        homeworkOrExamPoolDAO = HomeworkOrExamPoolDAO(db_util)
        entity_list = homeworkOrExamPoolDAO.query(subject, "courseName", is_all=True)
        result_list = []
        hepID_list = []
        if entity_list != "error":
            for i in entity_list:
                hepID_list.append(i.hepID)
            kp_results: tuple = homeworkOrExamPoolDAO.query_kp_by_hepIDs(hepID_list)
            if len(kp_results) == 0:
                return False
            n = 0
            for i in entity_list:
                result: dict = json.loads(i.question)
                result["answer"] = json.loads(i.answer)
                result["subject"] = subject
                result["difficulty"] = i.difficultyLevel
                result["knowledge_point"] = []
                while n < len(kp_results):
                    if kp_results[n][0] == i.hepID:
                        result["knowledge_point"].append(kp_results[n][1])
                        n += 1
                    else:
                        break
                result_list.append(result)
            with open(path, "w", encoding="utf-8") as f:
                f.write(json.dumps(result_list, ensure_ascii=False))
            return True
        else:
            return False

    @staticmethod
    def _readOurJson(json_path: str):
        with open(json_path, "r", encoding="utf-8") as f:
            t = json.load(f)
        return t

    @staticmethod
    def _query_diff_to_kp(type: str, diff_level: list, db_util):
        homeworkOrExamPoolDAO = HomeworkOrExamPoolDAO(db_util)
        values = (type,)
        placeholders = []
        results_dict = dict()
        for i in diff_level:
            placeholders.append("difficultyLevel = " + str(i))
            results_dict[i] = list()
        if len(placeholders) > 1:
            placeholders = " or ".join(placeholders)
        else:
            placeholders = placeholders[0]
        query = f"select difficultyLevel, hep_and_kp_mediater.hepID, kpName from homework_or_exam_pool, hep_and_kp_mediater where type = %s and homework_or_exam_pool.hepID = hep_and_kp_mediater.hepID and ({placeholders}) order by difficultyLevel"
        results = homeworkOrExamPoolDAO.execute_query(query, values)
        temp_n = 0
        # 这个循环里最后一个元素会丢失
        while temp_n < len(results) - 1:
            temp_n2 = 0
            n2 = results[temp_n][0]
            while n2 == results[temp_n][0] and temp_n < len(results) - 1:
                n = results[temp_n][1]
                results_dict[results[temp_n][0]].append(list())
                results_dict[results[temp_n][0]][temp_n2].append(n)
                while results[temp_n][1] == n:
                    results_dict[results[temp_n][0]][temp_n2].append(results[temp_n][2])
                    temp_n += 1
                    if temp_n > len(results) - 1:
                        temp_n -= 1
                        break
                temp_n2 += 1
        # 单独处理results的最后一个元素
        if len(results) > 0:
            end_element = results[-1]
            if len(results_dict[end_element[0]]) > 0:
                results_dict[end_element[0]][-1].append(end_element[2])
            else:
                results_dict[end_element[0]].append(list())
                results_dict[end_element[0]][-1].append(end_element[1])
                results_dict[end_element[0]][-1].append(end_element[2])
            del end_element
        # 检查是否存在某个难度的题目数量为0
        for i in results_dict:
            if len(results_dict[i]) == 0:
                return i
            # else:
            #     random.shuffle(results_dict[i])
        return results_dict

    @staticmethod
    def _analysis_diff_to_kp(diff_dict0: dict, kp_dict0: dict, query_dict0: dict):
        diff_dict = copy.deepcopy(diff_dict0)
        kp_dict = copy.deepcopy(kp_dict0)
        query_dict = copy.deepcopy(query_dict0)
        result_list = list()
        kp_list_result = list()
        executing_flag = 0

        for diff in diff_dict:
            random.shuffle(query_dict[diff])

            # 初始化
            kp_list = list(kp_dict.keys())
            random.shuffle(kp_list)
            kp_name_index = random.choice(list(range(len(kp_list))))
            temp_kp_name_index = []

            while True:
                # 循环停止的条件：该难度的题目数量已被消耗完，或者所有非零知识点都轮完
                del_index, find_flag, break_flag = [], 0, 0
                for query_index in range(len(query_dict[diff])):
                    # 还需处理最后一次出现的异常
                    if len(kp_list) == 0:
                        return [kp_dict.keys(), None]
                    # 开始寻找
                    if kp_list[kp_name_index] in query_dict[diff][query_index][1:]:
                        kp_list_result.append(query_dict[diff][query_index][1:])
                        kp_dict[kp_list[kp_name_index]] -= 1
                        diff_dict[diff] -= 1
                        del_index.append(query_index)
                        result_list.append(query_dict[diff][query_index][0])
                        if kp_dict[kp_list[kp_name_index]] <= 0 and diff_dict[diff] <= 0:
                            del kp_dict[kp_list[kp_name_index]]
                            break_flag = 1  # 循环可以停止了
                            break
                        elif kp_dict[kp_list[kp_name_index]] <= 0 and diff_dict[diff] > 0:
                            temp_kp_name_index.append(kp_list[kp_name_index])
                            del kp_dict[kp_list[kp_name_index]]  # 此时应该轮换到下一个知识点
                            del kp_list[kp_name_index]
                            if len(kp_list) > 0:
                                kp_name_index = random.choice(list(range(len(kp_list))))
                            break
                        elif kp_dict[kp_list[kp_name_index]] > 0 and diff_dict[diff] <= 0:
                            break_flag = 1  # 循环可以停止了
                            break
                        find_flag += 1
                    else:
                        find_flag += 1

                # 检查是否应该剔除查询结果：
                if len(del_index) > 0:
                    m = 0
                    for _ in del_index:
                        del query_dict[diff][_ - m]
                        m += 1

                # 检查是否可以break
                if break_flag == 1:
                    break

                # 此时该难度题目任有余，但是却没有与知识点匹配的查询结果，将此次的知识点冻结，换到下一个知识点
                if find_flag == len(query_dict[diff]) and kp_dict[kp_list[kp_name_index]] > 0:
                    temp_kp_name_index.append(kp_list[kp_name_index])
                    del kp_list[kp_name_index]
                    if len(kp_list) > 0:
                        kp_name_index = random.choice(list(range(len(kp_list))))

        print(result_list, kp_list_result)
        return [result_list, kp_list_result]

    @staticmethod
    def _analysis_paper(paper, diff_list, type_list, kp_list):
        """
        输入存储的试卷dict，难度组成的列表，类型组成的列表，知识点组成的列表\n
        输出文本分析和提取后的详细数据
        :param paper:
        :param diff_list:
        :param type_list:
        :param kp_list:
        :return:
        """
        # 获取总分
        sum_score = 0
        for _ in paper["score"]:
            for __ in _:
                for ___ in __:
                    sum_score += ___
        # 构建每道大题分数的字典，并分析分数占比
        score_dict = dict()
        score_text = f"本试卷总分：{sum_score}分,其中"
        for i in range(len(type_list)):
            s = 0
            score_dict[type_list[i]] = list()
            for j in paper["score"][i]:
                for k in j:
                    s += k
            score_dict[type_list[i]].append(s)
            score_text += f"{type_list[i]}一共{s}分,"
            score_dict[type_list[i]].append(s/sum_score)
            score_text += f'占比{round(s/sum_score*100, 2)}%;'
        score_text += "\n"

        # 获取平均难度
        ave_diff = round(sum(diff_list)/len(diff_list), 2)
        # 分析难度占比
        diff_text = f"该试卷平均难度为{ave_diff},"
        if 4 < ave_diff < 6:
            diff_text += "难度适中"
        elif 2 < ave_diff <= 4:
            diff_text += "难度偏易"
        elif 0 < ave_diff <= 2:
            diff_text += "难度简单"
        elif 6 <= ave_diff < 8:
            diff_text += "难度偏难"
        else:
            diff_text += "难度困难"
        diff_text += "\n"

        # 分析知识点情况

        return score_text + diff_text


    @staticmethod
    def _markdown_to_word(md_text, output_path):
        def add_text_to_paragraph(paragraph, element):
            for sub_element in element.children:
                run = paragraph.add_run(sub_element.get_text())
                if sub_element.name == 'strong':
                    run.bold = True
                if sub_element.name == 'em':
                    run.italic = True
                # 设置字体为宋体
                run.font.name = '宋体'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        def add_html_to_docx(html_content, doc):
            soup = BeautifulSoup(html_content, "html.parser")

            for element in soup.descendants:
                if element.name == 'h1':
                    paragraph = doc.add_heading(level=1)
                    run = paragraph.add_run(element.get_text())
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run.font.size = Pt(24)
                elif element.name == 'h2':
                    doc.add_heading(element.get_text(), level=2)
                elif element.name == 'h3':
                    doc.add_heading(element.get_text(), level=3)
                elif element.name == 'p':
                    paragraph = doc.add_paragraph()
                    add_text_to_paragraph(paragraph, element)
                elif element.name == 'li':
                    paragraph = doc.add_paragraph(style='ListBullet')
                    add_text_to_paragraph(paragraph, element)
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        paragraph = doc.add_paragraph(style='ListBullet')
                        add_text_to_paragraph(paragraph, li)

        html_content = markdown2.markdown(md_text, extras=["tables", "fenced-code-blocks"])

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        add_html_to_docx(html_content, doc)

        doc.save(output_path)
